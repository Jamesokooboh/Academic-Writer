import io
import uuid

import docx
import pytest
from fastapi.testclient import TestClient

from app.db.models import Chunk, Document, Sentence, SentenceStatus, User
from app.db.session import SessionLocal
from app.domain.document_export import build_final_text, to_docx, to_markdown, to_pdf
from app.domain.document_import import extract_text
from app.main import app


def test_extract_text_from_docx():
    document = docx.Document()
    document.add_paragraph("First paragraph.")
    document.add_paragraph("Second paragraph.")
    buffer = io.BytesIO()
    document.save(buffer)

    text = extract_text("essay.docx", buffer.getvalue())
    assert text == "First paragraph.\n\nSecond paragraph."


def test_extract_text_from_plain_and_markdown():
    content = "Paragraph one.\n\nParagraph two.".encode("utf-8")
    assert extract_text("notes.md", content) == "Paragraph one.\n\nParagraph two."
    assert extract_text("notes.txt", content) == "Paragraph one.\n\nParagraph two."


def test_extract_text_from_pdf_round_trips_through_our_own_export():
    pdf_bytes = to_pdf("A sentence about research findings.\n\nA second paragraph here.")
    text = extract_text("paper.pdf", pdf_bytes)
    assert "research findings" in text
    assert "second paragraph" in text


def test_to_markdown_is_utf8_passthrough():
    assert to_markdown("Hello world.") == b"Hello world."


def test_to_docx_round_trips_paragraphs():
    body = to_docx("Paragraph one.\n\nParagraph two.")
    reopened = docx.Document(io.BytesIO(body))
    paragraphs = [p.text for p in reopened.paragraphs if p.text]
    assert paragraphs == ["Paragraph one.", "Paragraph two."]


def _make_document_with_sentences(db) -> Document:
    user = User(email=f"test-{uuid.uuid4()}@example.com", password_hash="x")
    db.add(user)
    db.flush()
    document = Document(user_id=user.id, title="Export Test", writing_mode="Professional", word_count_mode="Balanced", rewrite_strength="Very Conservative")
    db.add(document)
    db.flush()

    chunk = Chunk(document_id=document.id, order_index=0, raw_text="orig1. orig2.", context_tail=None)
    db.add(chunk)
    db.flush()
    db.add(Sentence(chunk_id=chunk.id, order_index=0, original_text="orig1.", rewritten_text="REWRITTEN1.", status=SentenceStatus.REWRITTEN))
    db.add(Sentence(chunk_id=chunk.id, order_index=1, original_text="orig2.", status=SentenceStatus.GOOD))
    db.commit()
    db.refresh(document)
    return document


def test_build_final_text_uses_rewritten_text_only_when_status_rewritten():
    db = SessionLocal()
    try:
        document = _make_document_with_sentences(db)
        text = build_final_text(document)
        assert text == "REWRITTEN1. orig2."
    finally:
        db.close()


@pytest.fixture
def auth_client():
    with TestClient(app) as client:
        login = client.post("/api/auth/login", json={"email": "admin@example.com", "password": "test-password"})
        token = login.json()["access_token"]
        client.headers.update({"Authorization": f"Bearer {token}"})
        yield client


def test_import_then_export_round_trip(auth_client):
    source = docx.Document()
    source.add_paragraph("Imported paragraph one.")
    source.add_paragraph("Imported paragraph two.")
    buffer = io.BytesIO()
    source.save(buffer)

    import_response = auth_client.post(
        "/api/documents/import",
        files={"file": ("essay.docx", buffer.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert import_response.status_code == 201
    document_id = import_response.json()["id"]
    assert import_response.json()["title"] == "essay"

    analyze = auth_client.post(f"/api/documents/{document_id}/analyze")
    assert analyze.status_code == 200

    md_export = auth_client.get(f"/api/documents/{document_id}/export", params={"format": "md"})
    assert md_export.status_code == 200
    assert md_export.headers["content-type"].startswith("text/markdown")
    assert b"Imported paragraph one." in md_export.content

    docx_export = auth_client.get(f"/api/documents/{document_id}/export", params={"format": "docx"})
    assert docx_export.status_code == 200
    reopened = docx.Document(io.BytesIO(docx_export.content))
    assert any("Imported paragraph one" in p.text for p in reopened.paragraphs)

    pdf_export = auth_client.get(f"/api/documents/{document_id}/export", params={"format": "pdf"})
    assert pdf_export.status_code == 200
    assert pdf_export.headers["content-type"] == "application/pdf"
    assert pdf_export.content.startswith(b"%PDF")


def test_export_rejects_unknown_format(auth_client):
    create = auth_client.post("/api/documents", json={"title": "Doc", "content": "Some text."})
    document_id = create.json()["id"]
    response = auth_client.get(f"/api/documents/{document_id}/export", params={"format": "epub"})
    assert response.status_code == 400


def test_import_rejects_empty_file(auth_client):
    response = auth_client.post(
        "/api/documents/import", files={"file": ("empty.txt", b"   \n  ", "text/plain")}
    )
    assert response.status_code == 400
